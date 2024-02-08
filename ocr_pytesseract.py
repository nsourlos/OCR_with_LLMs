#Base Prompt: Given a folder in desktop get each image on it, and if it has text print it in the format  'name_of_image','text'.
# Then, rotate the image 90 degrees and repeat that. Keep rotating until the image is rotated 270 degrees. 
# After all these rotations detect the part(s) of each image that contain text. 
# Extract those parts and create a temporary image with only those part resized to the resolution of the original image. 
# Repeat the process of text recognition and rotations for the temporarily image and print the results. The folder path will be set in terminal command. 
#Also print the name of the image being processed first and then a progress bar
#Then give error in prompt (ValueError: All strings must be XML compatible: Unicode or ASCII, no NULL bytes or control characters) 
# and modify add_image_to_doc with the sanitized_text argument
#Then add image in doc.add_picture from cv2.imread
#Save word in Desktop
#When adding image to word do not crop it to fit the original dimensions but rather adapt to its new size
###return the rotated image without black around it with all the information that exists in the original image

#sudo apt-get update
#sudo apt-get upgrade
#sudo apt-get install libjpeg-dev libtiff-dev zlib1g-dev libfreetype6-dev liblcms2-dev libwebp-dev tcl8.6-dev tk8.6-dev python3-tk
#sudo apt install tesseract-ocr

##which pytesseract
##export PATH=/path/to/tesseract:$PATH
##source ~/.bashrc   # or ~/.bash_profile or ~/.zshrc, depending on your shell
##tesseract --version

#pip install python-docx==0.8.11 Pillow==8.1.2 opencv-python==4.9.0.80 tqdm==4.66.1 #pytesseract==0.3.7
#python script.py /path/to/image/folder

import os
import cv2
import pytesseract
from PIL import Image
from tqdm import tqdm
from docx import Document
from docx.shared import Inches
from io import BytesIO
import time

#import numpy as np

def rotate_image(image, degrees):
    # Rotate the image by the specified degrees
    rows, cols, _ = image.shape
    M = cv2.getRotationMatrix2D((cols / 2, rows / 2), degrees, 1)
    #rotated_image = cv2.warpAffine(image, M, (int(cols * 1.5), int(rows * 1.5)), flags=cv2.INTER_LINEAR) #To avoid cutting parts of the image
    rotated_image = cv2.warpAffine(image, M, (cols,rows), flags=cv2.INTER_LINEAR)     
        
    #cos_theta = np.abs(M[0, 0])
    #sin_theta = np.abs(M[0, 1])

    #new_width = int((rows * sin_theta) + (cols * cos_theta))
    #new_height = int((rows * cos_theta) + (cols * sin_theta))

    #M[0, 2] += (new_width / 2) - cols / 2
    #M[1, 2] += (new_height / 2) - rows / 2

    ## Apply the updated affine transformation
    #rotated_image = cv2.warpAffine(image, M, (new_width, new_height), flags=cv2.INTER_LINEAR)

    return rotated_image

def resize_image(image, scale_factor):
    # Resize the image by maintaining the original aspect ratio
    width = int(image.shape[1] * scale_factor)
    height = int(image.shape[0] * scale_factor)
    return cv2.resize(image, (width, height), interpolation=cv2.INTER_LINEAR)

def extract_text(image):
    # Use pytesseract to extract text from the image
    return pytesseract.image_to_string(image)
    
def add_image_to_doc(doc, image_path, name, text, rotated_image,rotation):
    # Replace problematic characters in the text
    sanitized_text = "".join(c if c.isprintable() else " " for c in text)

    # Add a paragraph with image name and sanitized text
    doc.add_paragraph(f"Image Name: {name}, Rotation: {rotation} degrees")
    doc.add_paragraph(f"Text: {sanitized_text}")
    
    # Calculate the scale factor for resizing
    scale_factor = 1 / 5.0  # Approximately 5 times less

    # Resize the rotated image
    resized_image = resize_image(rotated_image, scale_factor)

    # Convert OpenCV image to PIL image
    pil_image = Image.fromarray(cv2.cvtColor(resized_image, cv2.COLOR_BGR2RGB))

    # Save PIL image to BytesIO buffer
    image_buffer = BytesIO()
    pil_image.save(image_buffer, format="PNG")  # You can adjust the format as needed

    # Add the image to the document
    doc.add_picture(image_buffer, width=Inches(2.0))  # Adjust the width as needed

    # Add a line break between images
    doc.add_paragraph()

def main(folder_path):
    # Ensure the folder path exists
    if not os.path.exists(folder_path):
        print(f"Folder '{folder_path}' does not exist.")
        return

    # Get the list of image files in the folder
    image_files = [filename for filename in os.listdir(folder_path) if filename.endswith(('.jpg', '.jpeg', '.png'))]
    
    # Create a new Word document
    doc = Document()
    
    folder_name = os.path.basename(folder_path.rstrip(os.path.sep))
    output_filename = os.path.join(folder_path, f"{folder_name}_word.docx")

    # Iterate through each image file
    for filename in tqdm(image_files, desc="Processing Images and Rotations"): 
        print(filename, 'is being processed....')
        image_path = os.path.join(folder_path, filename)
        
        # Read the original image
        original_image = cv2.imread(image_path)

        # Rotate the image 90 degrees at a time
        for rotation in range(0, 271, 90):
            # if rotation==0 or rotation==270: #90, 180 give jebberish text (it happens only if original image correctly aligned)
                rotated_image = rotate_image(original_image, rotation)

                # Extract text from the rotated image
                text = extract_text(rotated_image)
                
                # Add image, name, and text to the Word document
                add_image_to_doc(doc, image_path, filename, text,rotated_image,rotation)

                # Print the image name, rotation, and extracted text
                print(f"{filename}, {rotation} degrees ")

                # Save the rotated image temporarily for text region extraction
                temp_image_path = f"temp_{rotation}_{filename}"
                cv2.imwrite(temp_image_path, rotated_image)

                # Open the temporarily saved image for text region extraction
                temp_image = cv2.imread(temp_image_path)

                # Extract only the text region from the temporary image
                text_region = extract_text(temp_image)

                # Remove the temporary image
                os.remove(temp_image_path)
            
    # Save the Word document with the folder name on the desktop
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    output_filename = os.path.join(desktop_path, f"{folder_name}_word.docx")
    doc.save(output_filename)
    print(f"Word document saved to {output_filename}")


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python text_extraction_script.py <folder_path>")
    else:
        folder_path = sys.argv[1]
        start=time.time()
        main(folder_path)
        end=time.time()
        print("It took", end-start, 'second to run')